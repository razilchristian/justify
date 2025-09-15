from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from openai import OpenAI
import os
import json
import tempfile
from docx import Document
from datetime import datetime
import PyPDF2
import re
import time
import pandas as pd
import numpy as np
import sqlite3
from contextlib import contextmanager
import logging
from logging.handlers import RotatingFileHandler
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import joblib

app = Flask(__name__)
CORS(app)

# Initialize rate limiting
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["200 per day", "50 per hour"]
)

# Initialize OpenAI (you'll need to set your API key in environment variables)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None

# Set up logging
if not app.debug:
    file_handler = RotatingFileHandler('legal_assistant.log', maxBytes=10240, backupCount=10)
    file_handler.setFormatter(logging.Formatter(
        '%(asctime)s %(levelname)s: %(message)s [in %(pathname)s:%(lineno)d]'
    ))
    file_handler.setLevel(logging.INFO)
    app.logger.addHandler(file_handler)
    app.logger.setLevel(logging.INFO)
    app.logger.info('Legal Assistant startup')

# Database setup
def get_db():
    conn = sqlite3.connect('legal_assistant.db')
    conn.row_factory = sqlite3.Row
    return conn

@contextmanager
def db_connection():
    conn = get_db()
    try:
        yield conn
    finally:
        conn.close()

def init_db():
    with db_connection() as conn:
        conn.execute('''
            CREATE TABLE IF NOT EXISTS chat_sessions (
                id TEXT PRIMARY KEY,
                history TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS uploaded_documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id TEXT,
                text TEXT,
                filename TEXT,
                file_type TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (session_id) REFERENCES chat_sessions (id)
            )
        ''')
        conn.execute('''
            CREATE TABLE IF NOT EXISTS api_usage (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                endpoint TEXT,
                ip_address TEXT,
                usage_count INTEGER DEFAULT 1,
                last_used TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        conn.commit()

# Initialize database
init_db()

# System prompt for legal assistant
LEGAL_ASSISTANT_PROMPT = """
You are JustiBot, an AI legal assistant specialized in Indian law. You provide accurate, helpful legal information while always reminding users that you are an AI and not a substitute for professional legal advice.

Key capabilities:
1. Legal drafting (contracts, petitions, notices)
2. Case law research and analysis
3. Legal document review
4. Explanation of legal concepts
5. Scenario analysis and issue spotting
6. Compliance guidance

Always:
- Be precise and cite relevant Indian laws when possible
- Maintain professional tone
- Suggest consulting actual lawyers for serious matters
- Acknowledge limitations when unsure
"""

# Enhanced legal knowledge base for fallback responses
LEGAL_KNOWLEDGE_BASE = {
    "contracts": {
        "elements": "Essential elements of a valid contract under Indian Contract Act, 1872: 1. Offer and Acceptance, 2. Lawful Consideration, 3. Capacity of Parties, 4. Free Consent, 5. Lawful Object, 6. Certainty and Possibility of Performance.",
        "types": "Common contract types in India: Service Agreement, Sale Deed, Lease Agreement, Employment Contract, Partnership Deed, Non-Disclosure Agreement (NDA), Memorandum of Understanding (MoU).",
        "breach": "Remedies for breach of contract: 1. Damages (Compensatory, Nominal, Liquidated), 2. Specific Performance, 3. Injunction, 4. Quantum Meruit, 5. Rescission."
    },
    "constitution": {
        "fundamental_rights": "Fundamental Rights under Indian Constitution (Part III): 1. Right to Equality (Articles 14-18), 2. Right to Freedom (Articles 19-22), 3. Right against Exploitation (Articles 23-24), 4. Right to Freedom of Religion (Articles 25-28), 5. Cultural and Educational Rights (Articles 29-30), 6. Right to Constitutional Remedies (Article 32).",
        "directive_principles": "Directive Principles of State Policy (Part IV) are guidelines for governance, though not enforceable in courts. They aim to establish social and economic democracy."
    },
    "ipc": {
        "common_sections": "Important IPC Sections: 302 (Murder), 304 (Culpable Homicide), 375 (Rape), 378 (Theft), 420 (Cheating), 499 (Defamation), 497 (Adultery - struck down in 2018).",
        "bailable": "Bailable offenses generally include less serious crimes where bail can be claimed as a right. Non-bailable offenses are more serious where bail is at court's discretion."
    },
    "civil_procedure": {
        "stages": "Stages of civil suit: 1. Institution of suit, 2. Issue of summons, 3. Appearance of parties, 4. Written statement, 5. Discovery and inspection, 6. Framing of issues, 7. Evidence, 8. Arguments, 9. Judgment, 10. Appeal.",
        "limitation": "The Limitation Act, 1963 prescribes time limits for different legal actions. For example: 3 years for contract disputes, 1 year for tort of defamation."
    },
    "consumer_protection": {
        "rights": "Consumer Rights under Consumer Protection Act, 2019: 1. Right to Safety, 2. Right to Information, 3. Right to Choose, 4. Right to be Heard, 5. Right to Redressal, 6. Right to Consumer Education.",
        "forums": "Consumer forums: District Commission (claims up to ₹1 crore), State Commission (claims ₹1 crore to ₹10 crores), National Commission (claims above ₹10 crores)."
    }
}

# Legal database and training data
class LegalDatabase:
    def __init__(self):
        self.cases = []
        self.statutes = []
        self.training_data = []
        self.enhanced_cases = self.load_enhanced_cases()
        self.vectorizer = None
        self.justice_dataset = None
        self.tfidf_matrix = None
    
    def load_justice_dataset(self):
        """Load and process the justice.csv dataset"""
        try:
            justice_path = "legal_datasets/justice.csv"
            if os.path.exists(justice_path):
                self.justice_dataset = pd.read_csv(justice_path)
                app.logger.info(f"Loaded justice dataset with {len(self.justice_dataset)} entries")
                
                # Preprocess the data for similarity search
                if 'question' in self.justice_dataset.columns and 'answer' in self.justice_dataset.columns:
                    # Combine question and answer for better search
                    self.justice_dataset['text'] = self.justice_dataset['question'] + " " + self.justice_dataset['answer']
                    
                    # Create TF-IDF vectorizer
                    self.vectorizer = TfidfVectorizer(stop_words='english', max_features=5000)
                    self.tfidf_matrix = self.vectorizer.fit_transform(self.justice_dataset['text'].fillna(''))
                    
                    # Save the vectorizer for later use
                    joblib.dump(self.vectorizer, 'models/tfidf_vectorizer.joblib')
                    joblib.dump(self.tfidf_matrix, 'models/tfidf_matrix.joblib')
                    joblib.dump(self.justice_dataset, 'models/justice_dataset.joblib')
                    
                    app.logger.info("Justice dataset processed and vectorized successfully")
                else:
                    app.logger.warning("Justice dataset doesn't have expected columns 'question' and 'answer'")
            else:
                app.logger.warning("Justice dataset not found at legal_datasets/justice.csv")
                # Create sample dataset if not exists
                os.makedirs("legal_datasets", exist_ok=True)
                sample_data = pd.DataFrame({
                    'question': [
                        'What is the basic structure doctrine?',
                        'How to file a consumer complaint?',
                        'What are fundamental rights?'
                    ],
                    'answer': [
                        'The basic structure doctrine was established in Kesavananda Bharati case, which holds that Parliament cannot amend the basic structure of the Constitution.',
                        'Consumer complaints can be filed with district, state, or national consumer forums based on the value of the claim.',
                        'Fundamental Rights are basic human rights enshrined in Part III of the Indian Constitution, including Right to Equality, Freedom, etc.'
                    ]
                })
                sample_data.to_csv(justice_path, index=False)
                self.justice_dataset = sample_data
                app.logger.info("Created sample justice dataset")
                
        except Exception as e:
            app.logger.error(f"Error loading justice dataset: {e}")
    
    def find_similar_questions(self, query, top_n=3):
        """Find similar questions in the justice dataset"""
        if self.vectorizer is None or self.tfidf_matrix is None:
            return []
        
        try:
            # Vectorize the query
            query_vec = self.vectorizer.transform([query])
            
            # Calculate cosine similarity
            similarities = cosine_similarity(query_vec, self.tfidf_matrix).flatten()
            
            # Get top N most similar questions
            top_indices = similarities.argsort()[-top_n:][::-1]
            
            results = []
            for idx in top_indices:
                if similarities[idx] > 0.1:  # Minimum similarity threshold
                    results.append({
                        'question': self.justice_dataset.iloc[idx]['question'],
                        'answer': self.justice_dataset.iloc[idx]['answer'],
                        'similarity': float(similarities[idx])
                    })
            
            return results
        except Exception as e:
            app.logger.error(f"Error finding similar questions: {e}")
            return []
    
    def load_enhanced_cases(self):
        """Load enhanced case information for better fallback responses"""
        return [
            {
                "title": "Kesavananda Bharati vs State of Kerala (1973)",
                "citation": "AIR 1973 SC 1461",
                "significance": "Established the Basic Structure Doctrine - Parliament cannot amend the basic structure of the Constitution",
                "key_points": ["Basic structure doctrine", "Judicial review of amendments", "Limitation on parliamentary power"]
            },
            {
                "title": "Minerva Mills Ltd. vs Union of India (1980)",
                "citation": "AIR 1980 SC 1789",
                "significance": "Strengthened basic structure doctrine, struck down parts of 42nd Amendment",
                "key_points": ["Balance between Fundamental Rights and DPSP", "Judicial review protection"]
            },
            {
                "title": "Maneka Gandhi vs Union of India (1978)",
                "citation": "AIR 1978 SC 597",
                "significance": "Expanded scope of Article 21 (Right to Life and Personal Liberty)",
                "key_points": ["Due process", "Procedure must be fair, just and reasonable", "Expanded personal liberty"]
            }
        ]
    
    def load_cases_from_csv(self, filepath):
        try:
            # Create data directory if it doesn't exist
            os.makedirs(os.path.dirname(filepath), exist_ok=True)
            
            # Check if file exists, if not create a sample one
            if not os.path.exists(filepath):
                sample_cases = pd.DataFrame([
                    {
                        "case_no": "AIR 1973 SC 1461",
                        "title": "Kesavananda Bharati vs State of Kerala",
                        "year": 1973,
                        "court": "Supreme Court of India",
                        "summary": "Landmark case that established the Basic Structure Doctrine of the Constitution",
                        "key_issues": "Constitutional amendments, Fundamental rights, Basic structure doctrine",
                        "outcome": "The Supreme Court outlined the basic structure doctrine of the Constitution"
                    },
                    {
                        "case_no": "AIR 1980 SC 1789",
                        "title": "Minerva Mills Ltd. vs Union of India",
                        "year": 1980,
                        "court": "Supreme Court of India",
                        "summary": "Strengthened the basic structure doctrine laid down in Kesavananda Bharati case",
                        "key_issues": "Constitutional amendments, Judicial review, Fundamental rights",
                        "outcome": "Struck down parts of the 42nd Amendment that prevented judicial review of constitutional amendments"
                    }
                ])
                sample_cases.to_csv(filepath, index=False)
                self.cases = sample_cases.to_dict('records')
            else:
                df = pd.read_csv(filepath)
                self.cases = df.to_dict('records')
            return len(self.cases)
        except Exception as e:
            app.logger.error(f"Error loading cases: {e}")
            # Return dummy data if file not found
            self.cases = [
                {
                    "case_no": "AIR 1973 SC 1461",
                    "title": "Kesavananda Bharati vs State of Kerala",
                    "year": 1973,
                    "court": "Supreme Court of India",
                    "summary": "Landmark case that established the Basic Structure Doctrine of the Constitution",
                    "key_issues": "Constitutional amendments, Fundamental rights, Basic structure doctrine",
                    "outcome": "The Supreme Court outlined the basic structure doctrine of the Constitution"
                },
                {
                    "case_no": "AIR 1980 SC 1789",
                    "title": "Minerva Mills Ltd. vs Union of India",
                    "year": 1980,
                    "court": "Supreme Court of India",
                    "summary": "Strengthened the basic structure doctrine laid down in Kesavananda Bharati case",
                    "key_issues": "Constitutional amendments, Judicial review, Fundamental rights",
                    "outcome": "Struck down parts of the 42nd Amendment that prevented judicial review of constitutional amendments"
                }
            ]
            return len(self.cases)
    
    def load_training_data(self):
        # Pre-load training examples for the AI
        self.training_data = [
            {
                "input": "How to draft a legal notice?",
                "output": "A legal notice should include: 1. sender and recipient details, 2. clear subject line, 3. facts of the case, 4. legal basis for the claim, 5. relief sought, 6. time given for compliance, and 7. consequences of non-compliance. Would you like me to help draft a specific notice?"
            },
            {
                "input": "What is the difference between civil and criminal law?",
                "output": "Civil law deals with disputes between individuals/organizations where compensation may be awarded. Criminal law deals with crimes against the state/society where punishment may be imposed. Key differences: purpose (compensation vs punishment), burden of proof (balance of probabilities vs beyond reasonable doubt), and initiating party (individual vs state)."
            },
            {
                "input": "How to file a consumer complaint?",
                "output": "To file a consumer complaint: 1. Gather all documents ( bills, correspondence, evidence), 2. Draft a complaint with facts and relief sought, 3. File with the appropriate consumer forum based on the value of claim, 4. Pay the required fee. The process is designed to be consumer-friendly and doesn't necessarily require a lawyer."
            }
        ]
        return len(self.training_data)

# Initialize legal database
legal_db = LegalDatabase()

def preload_training_data():
    """Pre-load training data and cases at startup"""
    app.logger.info("Loading legal cases...")
    num_cases = legal_db.load_cases_from_csv("data/legal_cases.csv")
    
    app.logger.info("Loading training data...")
    num_training = legal_db.load_training_data()
    
    app.logger.info("Loading justice dataset...")
    legal_db.load_justice_dataset()
    
    app.logger.info(f"Loaded {num_cases} legal cases, {num_training} training examples, and justice dataset")
    return num_cases, num_training

def track_api_usage(endpoint):
    """Track API usage for rate limiting and analytics"""
    ip_address = get_remote_address()
    with db_connection() as conn:
        # Check if this IP has used this endpoint recently
        cursor = conn.execute(
            "SELECT * FROM api_usage WHERE endpoint = ? AND ip_address = ?",
            (endpoint, ip_address)
        )
        existing = cursor.fetchone()
        
        if existing:
            # Update existing record
            conn.execute(
                "UPDATE api_usage SET usage_count = usage_count + 1, last_used = CURRENT_TIMESTAMP WHERE id = ?",
                (existing['id'],)
            )
        else:
            # Create new record
            conn.execute(
                "INSERT INTO api_usage (endpoint, ip_address) VALUES (?, ?)",
                (endpoint, ip_address)
            )
        conn.commit()

def extract_text_from_pdf(file_content):
    """Improved PDF text extraction function"""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name
        
        text = ""
        with open(tmp_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        
        os.unlink(tmp_path)
        return text
    except Exception as e:
        app.logger.error(f"PDF extraction error: {e}")
        return ""

def get_fallback_response(user_message):
    """Generate a fallback response when OpenAI API is unavailable"""
    user_message_lower = user_message.lower()
    
    # First try to find similar questions in justice dataset
    similar_questions = legal_db.find_similar_questions(user_message)
    if similar_questions:
        best_match = similar_questions[0]
        if best_match['similarity'] > 0.3:  # Good match threshold
            return f"{best_match['answer']}\n\n[Answer from legal knowledge base - Similarity: {best_match['similarity']:.2f}]"
    
    # Greeting responses
    if any(word in user_message_lower for word in ['hello', 'hi', 'hey', 'greetings']):
        return "Hello! I'm JustiBot, your AI legal assistant. I'm currently operating in limited mode. How can I help you with Indian legal information today?"
    
    # Contract-related queries
    elif any(word in user_message_lower for word in ['contract', 'draft', 'agreement']):
        return f"{LEGAL_KNOWLEDGE_BASE['contracts']['elements']}\n\n{LEGAL_KNOWLEDGE_BASE['contracts']['types']}"
    
    # Case law queries
    elif any(word in user_message_lower for word in ['case', 'judgment', 'precedent']):
        case_info = "\n\n".join([f"{case['title']} ({case['citation']}): {case['significance']}" 
                               for case in legal_db.enhanced_cases[:2]])
        return f"Here are some important Indian legal cases:\n\n{case_info}"
    
    # Fundamental rights queries
    elif any(word in user_message_lower for word in ['right', 'fundamental', 'constitution']):
        return LEGAL_KNOWLEDGE_BASE['constitution']['fundamental_rights']
    
    # Consumer protection queries
    elif any(word in user_message_lower for word in ['consumer', 'complaint', 'forum']):
        return f"{LEGAL_KNOWLEDGE_BASE['consumer_protection']['rights']}\n\n{LEGAL_KNOWLEDGE_BASE['consumer_protection']['forums']}"
    
    # IPC queries
    elif any(word in user_message_lower for word in ['ipc', 'section', 'penal']):
        return LEGAL_KNOWLEDGE_BASE['ipc']['common_sections']
    
    # General legal help
    else:
        return ("I'm currently operating in limited mode due to API restrictions. "
                "I can help with information about: Indian Contract Law, Constitution, "
                "IPC sections, Consumer Protection, and important legal cases. "
                "Please ask me specific legal questions about Indian law.")

@app.route('/')
def serve_index():
    try:
        return send_file('index.html')
    except:
        return "HTML file not found. Please make sure index.html is in the same directory."

@app.route('/openai/chat', methods=['POST'])
@limiter.limit("10 per minute")
def chat_with_openai():
    try:
        track_api_usage('/openai/chat')
        
        data = request.json
        user_message = data.get('message', '')
        history = data.get('history', [])
        session_id = data.get('session_id', 'default')
        model = data.get('model', 'gpt-3.5-turbo')
        
        # Save chat history to database
        with db_connection() as conn:
            conn.execute(
                "INSERT OR REPLACE INTO chat_sessions (id, history, updated_at) VALUES (?, ?, CURRENT_TIMESTAMP)",
                (session_id, json.dumps(history + [{"role": "user", "content": user_message}]))
            )
            conn.commit()
        
        # Prepare messages with system prompt and history
        messages = [{"role": "system", "content": LEGAL_ASSISTANT_PROMPT}]
        
        # Add history if provided
        for msg in history:
            messages.append({"role": msg["role"], "content": msg["content"]})
        
        # Add current message
        messages.append({"role": "user", "content": user_message})
        
        ai_response = ""
        
        # Try to call OpenAI API, use fallback if it fails
        if client:
            try:
                # Call OpenAI API with the new client format
                response = client.chat.completions.create(
                    model=model,
                    messages=messages,
                    temperature=0.2,
                    max_tokens=900
                )
                
                ai_response = response.choices[0].message.content
                
            except Exception as api_error:
                app.logger.error(f"OpenAI API error, using fallback: {api_error}")
                ai_response = get_fallback_response(user_message)
                ai_response += "\n\n[Note: Currently using fallback mode due to API limitations]"
        else:
            ai_response = get_fallback_response(user_message)
            ai_response += "\n\n[Note: Operating in fallback mode - OpenAI API not configured]"
        
        # Update chat history with AI response
        with db_connection() as conn:
            conn.execute(
                "UPDATE chat_sessions SET history = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (json.dumps(history + [
                    {"role": "user", "content": user_message},
                    {"role": "assistant", "content": ai_response}
                ]), session_id)
            )
            conn.commit()
        
        return jsonify({"response": ai_response})
    
    except Exception as e:
        app.logger.error(f"Chat error: {e}")
        # Return fallback response even for other errors
        fallback_response = get_fallback_response(user_message if 'user_message' in locals() else "")
        return jsonify({"response": fallback_response})

@app.route('/api/justice-search', methods=['POST'])
def justice_search():
    """Search for similar questions in the justice dataset"""
    try:
        data = request.json
        query = data.get('query', '')
        top_n = data.get('top_n', 5)
        
        similar_questions = legal_db.find_similar_questions(query, top_n)
        
        return jsonify({
            "query": query,
            "results": similar_questions,
            "total_results": len(similar_questions)
        })
    
    except Exception as e:
        app.logger.error(f"Justice search error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/api/justice-dataset', methods=['GET'])
def get_justice_dataset():
    """Get information about the justice dataset"""
    try:
        if legal_db.justice_dataset is None:
            return jsonify({"error": "Justice dataset not loaded"}), 404
        
        return jsonify({
            "total_entries": len(legal_db.justice_dataset),
            "columns": list(legal_db.justice_dataset.columns),
            "sample_questions": legal_db.justice_dataset['question'].head(10).tolist() if 'question' in legal_db.justice_dataset.columns else []
        })
    
    except Exception as e:
        app.logger.error(f"Justice dataset info error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/upload', methods=['POST'])
@limiter.limit("20 per hour")
def upload_document():
    try:
        track_api_usage('/upload')
        
        data = request.json
        text = data.get('text', '')
        session_id = data.get('session_id', 'default')
        filename = data.get('filename', 'document')
        file_type = data.get('file_type', 'text')
        
        # If it's a PDF, extract text
        if file_type == 'pdf' and text.startswith('data:application/pdf;base64,'):
            import base64
            base64_data = text.split(',', 1)[1]
            file_content = base64.b64decode(base64_data)
            text = extract_text_from_pdf(file_content)
        
        with db_connection() as conn:
            conn.execute(
                "INSERT INTO uploaded_documents (session_id, text, filename, file_type) VALUES (?, ?, ?, ?)",
                (session_id, text[:10000], filename, file_type)  # Limit text size
            )
            conn.commit()
        
        return jsonify({"status": "success", "message": "Document uploaded successfully"})
    
    except Exception as e:
        app.logger.error(f"Upload error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/analyze-documents', methods=['POST'])
@limiter.limit("15 per hour")
def analyze_documents():
    try:
        track_api_usage('/analyze-documents')
        
        if not client:
            return jsonify({"error": "OpenAI API key not configured"}), 500
            
        data = request.json
        session_id = data.get('session_id', 'default')
        question = data.get('question', '')
        
        # Get documents from database
        with db_connection() as conn:
            cursor = conn.execute(
                "SELECT text FROM uploaded_documents WHERE session_id = ? ORDER BY created_at DESC",
                (session_id,)
            )
            documents = cursor.fetchall()
        
        if not documents:
            return jsonify({"error": "No documents found for analysis"}), 400
        
        # Get all uploaded documents
        documents_text = "\n\n".join([doc['text'] for doc in documents])
        
        # Prepare analysis prompt
        analysis_prompt = f"""
        Analyze the following legal documents and answer the user's question.
        
        DOCUMENTS:
        {documents_text}
        
        QUESTION: {question}
        
        Provide a comprehensive analysis with:
        1. Key findings from the documents
        2. Relevant legal provisions
        3. Potential issues or concerns
        4. Recommendations
        """
        
        # Call OpenAI for analysis
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": LEGAL_ASSISTANT_PROMPT},
                {"role": "user", "content": analysis_prompt}
            ],
            temperature=0.2,
            max_tokens=1500
        )
        
        analysis = response.choices[0].message.content
        
        return jsonify({"analysis": analysis})
    
    except Exception as e:
        app.logger.error(f"Analysis error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/draft-document', methods=['POST'])
@limiter.limit("10 per hour")
def draft_document():
    try:
        track_api_usage('/draft-document')
        
        data = request.json
        doc_type = data.get('type', 'contract')
        requirements = data.get('requirements', '')
        jurisdiction = data.get('jurisdiction', 'India')
        
        # If OpenAI is not available, provide a template
        if not client:
            return jsonify({
                "draft": f"""# {doc_type.title()} Draft - Template
                
[This is a template as OpenAI service is currently unavailable]

**PARTIES:**
[Insert Party Names and Details]

**RECITALS:**
[Background and purpose of the {doc_type}]

**TERMS AND CONDITIONS:**

1. [Insert first term]
2. [Insert second term]
3. [Insert third term]

**GOVERNING LAW:**
This {doc_type} shall be governed by and construed in accordance with the laws of {jurisdiction}.

**SIGNATURES:**

_________________________
[Party A Name and Title]

_________________________
[Party B Name and Title]

**Note:** This is a template. Please consult with a qualified legal professional for specific advice."""
            })
        
        # Prepare drafting prompt
        draft_prompt = f"""
        Draft a {doc_type} for Indian jurisdiction with the following requirements:
        
        {requirements}
        
        Please provide a comprehensive, professionally formatted legal document with:
        1. Appropriate headings and sections
        2. Standard legal language
        3. Placeholders for specific details in [brackets]
        4. Relevant legal provisions for {jurisdiction}
        """
        
        # Call OpenAI for drafting
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": LEGAL_ASSISTANT_PROMPT},
                {"role": "user", "content": draft_prompt}
            ],
            temperature=0.3,
            max_tokens=2000
        )
        
        draft = response.choices[0].message.content
        
        return jsonify({"draft": draft})
    
    except Exception as e:
        app.logger.error(f"Drafting error: {e}")
        # Provide a basic template as fallback
        return jsonify({
            "draft": f"# Document Draft Template\n\nI encountered an error generating your document. Here's a basic template structure for a {data.get('type', 'contract')}:\n\n1. Parties Section\n2. Definitions\n3. Terms and Conditions\n4. Payment Terms (if applicable)\n5. Termination Clause\n6. Governing Law\n7. Signatures\n\nPlease try again later or consult a legal professional for specific drafting needs."
        })

@app.route('/find-judgments', methods=['POST'])
@limiter.limit("15 per hour")
def find_judgments():
    try:
        track_api_usage('/find-judgments')
        
        data = request.json
        issue = data.get('issue', '')
        jurisdiction = data.get('jurisdiction', 'Supreme Court of India')
        limit = data.get('limit', 5)
        
        # First, try to find matches in our database
        matching_cases = []
        for case in legal_db.cases:
            if (issue.lower() in case.get('key_issues', '').lower() or 
                issue.lower() in case.get('summary', '').lower() or
                issue.lower() in case.get('title', '').lower()):
                matching_cases.append(case)
                if len(matching_cases) >= limit:
                    break
        
        # If we found cases in our database, return them
        if matching_cases:
            return jsonify({
                "judgments": matching_cases,
                "source": "local_database"
            })
        
        # If no local matches and OpenAI is available, use it to find relevant judgments
        if client:
            judgment_prompt = f"""
            Based on your knowledge of Indian case law, provide information about relevant judgments for the following legal issue:
            
            {issue}
            
            Please include:
            1. Key case names and citations
            2. Summary of legal principles established
            3. Relevance to the issue presented
            4. Any limitations or subsequent developments
            """
            
            # Call OpenAI for judgment search
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": LEGAL_ASSISTANT_PROMPT},
                    {"role": "user", "content": judgment_prompt}
                ],
                temperature=0.2,
                max_tokens=1500
            )
            
            judgments = response.choices[0].message.content
            
            return jsonify({
                "judgments": judgments,
                "source": "openai"
            })
        else:
            return jsonify({
                "judgments": "OpenAI service not available. Please configure your API key.",
                "source": "error"
            })
    
    except Exception as e:
        app.logger.error(f"Judgment search error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/export-docx', methods=['POST'])
@limiter.limit("10 per hour")
def export_docx():
    try:
        track_api_usage('/export-docx')
        
        data = request.json
        content = data.get('content', '')
        filename = data.get('filename', 'legal_document')
        
        # Create a new Document
        doc = Document()
        
        # Add content to the document
        for line in content.split('\n'):
            if line.strip() == '':
                continue
            # Check if line looks like a heading
            if re.match(r'^[A-Z][A-Za-z\s]+:$', line) or re.match(r'^[IVX]+\.', line) or re.match(r'^[0-9]+\.', line):
                heading = doc.add_heading(line, level=2)
            else:
                paragraph = doc.add_paragraph(line)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return send_file(
            temp_file.name, 
            as_attachment=True, 
            download_name=f"{filename}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    except Exception as e:
        app.logger.error(f"Export error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/predict-outcome', methods=['POST'])
@limiter.limit("10 per hour")
def predict_outcome():
    try:
        track_api_usage('/predict-outcome')
        
        if not client:
            return jsonify({"error": "OpenAI API key not configured"}), 500
            
        data = request.json
        case_details = data.get('case_details', '')
        jurisdiction = data.get('jurisdiction', 'India')
        
        # Prepare prediction prompt
        prediction_prompt = f"""
        Based on the following case details, provide a predictive analysis of likely outcomes:
        
        {case_details}
        
        Please include:
        1. Assessment of strengths and weaknesses of the case
        2. Relevant legal precedents
        3. Potential arguments for each side
        4. Likely outcomes with probabilities
        5. Factors that could influence the outcome
        
        Remember to emphasize that this is predictive analysis only and not legal advice.
        """
        
        # Call OpenAI for prediction
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": LEGAL_ASSISTANT_PROMPT},
                {"role": "user", "content": prediction_prompt}
            ],
            temperature=0.3,
            max_tokens=1500
        )
        
        prediction = response.choices[0].message.content
        
        return jsonify({"prediction": prediction})
    
    except Exception as e:
        app.logger.error(f"Prediction error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/compliance-check', methods=['POST'])
@limiter.limit("10 per hour")
def compliance_check():
    try:
        track_api_usage('/compliance-check')
        
        if not client:
            return jsonify({"error": "OpenAI API key not configured"}), 500
            
        data = request.json
        document_text = data.get('document_text', '')
        regulations = data.get('regulations', 'Indian laws')
        
        # Prepare compliance check prompt
        compliance_prompt = f"""
        Review the following document for compliance with {regulations}:
        
        {document_text}
        
        Please provide:
        1. Identification of potential compliance issues
        2. Relevant legal requirements
        3. Recommendations for addressing any issues
        4. References to specific regulations or standards
        """
        
        # Call OpenAI for compliance check
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": LEGAL_ASSISTANT_PROMPT},
                {"role": "user", "content": compliance_prompt}
            ],
            temperature=0.2,
            max_tokens=1500
        )
        
        compliance_report = response.choices[0].message.content
        
        return jsonify({"compliance_report": compliance_report})
    
    except Exception as e:
        app.logger.error(f"Compliance check error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/due-diligence', methods=['POST'])
@limiter.limit("10 per hour")
def due_diligence():
    try:
        track_api_usage('/due-diligence')
        
        if not client:
            return jsonify({"error": "OpenAI API key not configured"}), 500
            
        data = request.json
        documents_text = data.get('documents_text', '')
        transaction_type = data.get('transaction_type', 'general')
        
        # Prepare due diligence prompt
        diligence_prompt = f"""
        Perform due diligence analysis on the following documents for a {transaction_type} transaction:
        
        {documents_text}
        
        Please provide:
        1. Identification of key risks and issues
        2. Legal implications of identified issues
        3. Recommendations for risk mitigation
        4. Priority areas for further investigation
        5. Potential impact on transaction structure
        """
        
        # Call OpenAI for due diligence
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": LEGAL_ASSISTANT_PROMPT},
                {"role": "user", "content": diligence_prompt}
            ],
            temperature=0.2,
            max_tokens=2000
        )
        
        diligence_report = response.choices[0].message.content
        
        return jsonify({"diligence_report": diligence_report})
    
    except Exception as e:
        app.logger.error(f"Due diligence error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/deposition-prep', methods=['POST'])
@limiter.limit("10 per hour")
def deposition_prep():
    try:
        track_api_usage('/deposition-prep')
        
        if not client:
            return jsonify({"error": "OpenAI API key not configured"}), 500
            
        data = request.json
        case_details = data.get('case_details', '')
        witness_role = data.get('witness_role', 'general')
        
        # Prepare deposition preparation prompt
        deposition_prompt = f"""
        Prepare deposition questions for a {witness_role} witness based on the following case details:
        
        {case_details}
        
        Please provide:
        1. Background and foundational questions
        2. Key factual questions specific to the case
        3. Questions to establish or challenge credibility
        4. Questions about documents or evidence
        5. Potential follow-up questions based on likely responses
        6. Strategy notes for the examining attorney
        """
        
        # Call OpenAI for deposition preparation
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": LEGAL_ASSISTANT_PROMPT},
                {"role": "user", "content": deposition_prompt}
            ],
            temperature=0.3,
            max_tokens=1800
        )
        
        deposition_questions = response.choices[0].message.content
        
        return jsonify({"deposition_questions": deposition_questions})
    
    except Exception as e:
        app.logger.error(f"Deposition prep error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        "status": "healthy",
        "uptime": time.time() - app.config.get("START_TIME", time.time()),
        "cases_loaded": len(legal_db.cases),
        "training_examples": len(legal_db.training_data),
        "justice_dataset_loaded": legal_db.justice_dataset is not None,
        "justice_dataset_size": len(legal_db.justice_dataset) if legal_db.justice_dataset is not None else 0,
        "openai_configured": bool(OPENAI_API_KEY),
        "openai_available": client is not None,
        "fallback_mode": client is None or not OPENAI_API_KEY
    })

@app.route('/sessions', methods=['GET'])
def get_sessions():
    try:
        with db_connection() as conn:
            cursor = conn.execute(
                "SELECT id, created_at, updated_at FROM chat_sessions ORDER BY updated_at DESC"
            )
            sessions = cursor.fetchall()
        
        return jsonify({"sessions": [dict(session) for session in sessions]})
    except Exception as e:
        app.logger.error(f"Session retrieval error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/sessions/<session_id>', methods=['GET'])
def get_session(session_id):
    try:
        with db_connection() as conn:
            cursor = conn.execute(
                "SELECT id, history, created_at, updated_at FROM chat_sessions WHERE id = ?",
                (session_id,)
            )
            session = cursor.fetchone()
            
            if not session:
                return jsonify({"error": "Session not found"}), 404
            
            cursor = conn.execute(
                "SELECT id, filename, file_type, created_at FROM uploaded_documents WHERE session_id = ? ORDER BY created_at DESC",
                (session_id,)
            )
            documents = cursor.fetchall()
        
        return jsonify({
            "session": dict(session),
            "documents": [dict(doc) for doc in documents]
        })
    except Exception as e:
        app.logger.error(f"Session retrieval error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route('/sessions/<session_id>', methods=['DELETE'])
def delete_session(session_id):
    try:
        with db_connection() as conn:
            conn.execute("DELETE FROM chat_sessions WHERE id = ?", (session_id,))
            conn.execute("DELETE FROM uploaded_documents WHERE session_id = ?", (session_id,))
            conn.commit()
        
        return jsonify({"status": "success", "message": "Session deleted successfully"})
    except Exception as e:
        app.logger.error(f"Session deletion error: {e}")
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    # Create necessary directories
    os.makedirs("data", exist_ok=True)
    os.makedirs("legal_datasets", exist_ok=True)
    os.makedirs("models", exist_ok=True)
    os.makedirs("backups", exist_ok=True)
    
    # Pre-load training data
    preload_training_data()
    
    app.config["START_TIME"] = time.time()
    app.logger.info("JustiBot AI Legal Assistant starting...")
    app.logger.info(f"OpenAI API Key: {'Loaded' if OPENAI_API_KEY else 'Missing'}")
    app.logger.info(f"OpenAI Client: {'Available' if client else 'Not available'}")
    app.logger.info(f"Loaded {len(legal_db.cases)} legal cases")
    app.logger.info(f"Pre-loaded {len(legal_db.training_data)} training examples")
    app.logger.info(f"Justice dataset: {'Loaded' if legal_db.justice_dataset is not None else 'Not available'}")
    
    if not client or not OPENAI_API_KEY:
        app.logger.warning("Running in fallback mode - OpenAI API not available")
    
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", 8080)), debug=True)